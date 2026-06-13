"""Generate vol2atlas_method.docx — 2-page method writeup.

Stage A — µCT → CCF
Stage B — EM → aligned µCT (already in CCF)

Run:  python scripts/make_method_doc.py
Out:  vol2atlas_method.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_run(run, *, font="Calibri", size=10, bold=False, italic=False,
            color=(0x22, 0x22, 0x22)):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, font="Calibri", size=18, bold=True, color=(0x1F, 0x3D, 0x7A))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, font="Calibri", size=12, bold=True, color=(0x1F, 0x3D, 0x7A))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


def lead(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=10, italic=True, color=(0x55, 0x55, 0x55))
    p.paragraph_format.space_after = Pt(6)


def para(doc, runs):
    """`runs` is a list of (text, kwargs) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    for txt, kw in runs:
        r = p.add_run(txt)
        set_run(r, size=kw.get("size", 10),
                bold=kw.get("bold", False),
                italic=kw.get("italic", False),
                font=kw.get("font", "Calibri"),
                color=kw.get("color", (0x22, 0x22, 0x22)))
    return p


def bullet(doc, runs):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    for txt, kw in runs:
        r = p.add_run(txt)
        set_run(r, size=kw.get("size", 10),
                bold=kw.get("bold", False),
                italic=kw.get("italic", False),
                font=kw.get("font", "Calibri"),
                color=kw.get("color", (0x22, 0x22, 0x22)))


def code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln)
        set_run(r, font="Consolas", size=9, color=(0x10, 0x40, 0x70))


def main():
    doc = Document()

    # Page setup — tighter margins to keep it to 2 pages
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    # ----- TITLE -----
    h1(doc, "vol2atlas — method writeup")
    lead(doc, "Status: Stage A (µCT → CCF) is implemented and tested on real "
              "µCT data. Stage B (EM → aligned µCT) is a design proposal — "
              "the --reference flag, alignFull --tps, and voxel-size-aware UI "
              "are unmerged, and no EM volume has been run through the "
              "pipeline. Read Stage B as the target architecture, not a "
              "shipping feature.")

    # ============================================================
    # STAGE A
    # ============================================================
    h2(doc, "Stage A  —  µCT  →  Allen CCF   [implemented, tested on µCT]")

    para(doc, [
        ("Input.  ", {"bold": True}),
        ("OME-Zarr µCT volume at µm-class voxels (GB scale). ", {}),
        ("Reference.  ", {"bold": True}),
        ("Allen CCF (allen_mouse_25um) loaded through BrainGlobe. ", {}),
        ("Output.  ", {"bold": True}),
        ("uct_in_ccf.zarr — multiscale OME-Zarr on the CCF voxel grid.", {}),
    ])

    para(doc, [("CLI.", {"bold": True})])
    code_block(doc, [
        "vol2atlas init      sample_uct.zarr -s state_uct.json --voxel-um 2.74",
        "vol2atlas prealign  state_uct.json     # 3D rough rigid + CCF crop ROI",
        "vol2atlas refine    state_uct.json     # ortho-view fine refinement",
        "vol2atlas landmarks state_uct.json     # optional Procrustes refit",
        "vol2atlas alignFull state_uct.json -o out/uct_in_ccf.zarr",
    ])

    para(doc, [("Method.", {"bold": True})])
    bullet(doc, [
        ("Display.  ", {"bold": True}),
        ("napari with a strided preview ≤ 192³ voxels — RAM stays ~30 MB regardless of source size.", {}),
    ])
    bullet(doc, [
        ("Rigid pose.  ", {"bold": True}),
        ("Six sliders (3 translations µm, 3 rotations °) plus three axis-flip checkboxes. "
         "The displayed sample is a live scipy resample onto the CCF voxel grid — WYSIWYG.", {}),
    ])
    bullet(doc, [
        ("Crop ROI.  ", {"bold": True}),
        ("Six min/max sliders crop the CCF to a bounding box around the sample. "
         "Subsequent steps work only on that subvolume.", {}),
    ])
    bullet(doc, [
        ("Refine.  ", {"bold": True}),
        ("Tighter ±3 mm / ±30° ranges in switchable axial / coronal / sagittal ortho views.", {}),
    ])
    bullet(doc, [
        ("Landmarks (optional).  ", {"bold": True}),
        ("Pick paired clicks on sample and CCF; Procrustes/Kabsch on ≥3 pairs "
         "replaces the saved rigid pose. Skip Fit if the hand pose already overlays correctly.", {}),
    ])
    bullet(doc, [
        ("Production warp.  ", {"bold": True}),
        ("alignFull applies the rigid transform to the full OME-Zarr at every pyramid level — "
         "one output chunk at a time via dask_image.ndinterp. RAM bounded by chunk × rotation overhead. "
         "TB-scale level 0 is supported as long as disk is.", {}),
    ])

    para(doc, [("Math.", {"bold": True})])
    para(doc, [
        ("All coordinates are (z, y, x) µm — no unit-tagging schema. "
         "Rigid: ", {}),
        ("output_µm = R · F · (input_µm − c) + c + t", {"font": "Consolas", "size": 9.5}),
        ("  with R the intrinsic ZYX Euler rotation, F = diag(±1, ±1, ±1) axis flips, "
         "c the sample center, t the translation. ", {}),
        ("Optional TPS (--tps): ", {"bold": True}),
        ("thin-plate spline on the residuals between the rigid prediction and landmark clicks. "
         "Vanishes far from landmarks → rigid result preserved.", {}),
    ])

    para(doc, [("State.", {"bold": True})])
    para(doc, [
        ("Every step reads and writes one state_uct.json:  ", {}),
        ("{sample_zarr, sample_level, sample_voxel_um, atlas_name, transform, landmarks, ccf_crop_bbox, history}.  ",
         {"font": "Consolas", "size": 9.5}),
        ("Resume any step.", {}),
    ])

    # ============================================================
    # STAGE B
    # ============================================================
    h2(doc, "Stage B  —  EM  →  aligned µCT   [ROADMAP — not implemented, not tested]")

    para(doc, [
        ("Status.  ", {"bold": True, "color": (0xB0, 0x30, 0x30)}),
        ("This section describes the target design. The --reference flag "
         "on init, alignFull --tps, the Reference abstraction, and the "
         "voxel-size-aware UI changes are all unmerged. No EM volume has "
         "been processed by this pipeline yet — claims about TB-scale "
         "behavior are extrapolated from Stage A's chunkwise µCT runs, "
         "not measured on EM.",
         {"color": (0xB0, 0x30, 0x30)}),
    ])

    para(doc, [
        ("Proposed input.  ", {"bold": True}),
        ("OME-Zarr EM volume at nm voxels (TB scale; stored as µm — 4 nm = 0.004 µm). ", {}),
        ("Proposed reference.  ", {"bold": True}),
        ("uct_in_ccf.zarr from Stage A. ", {}),
        ("Proposed output.  ", {"bold": True}),
        ("em_in_ccf.zarr on the same voxel grid as uct_in_ccf.zarr.", {}),
    ])

    para(doc, [("Proposed CLI (none of these flags exist today).", {"bold": True})])
    code_block(doc, [
        "vol2atlas init      sample_em.zarr -s state_em.json \\",
        "                    --reference out/uct_in_ccf.zarr --voxel-um 0.004",
        "vol2atlas prealign  state_em.json",
        "vol2atlas refine    state_em.json",
        "vol2atlas landmarks state_em.json",
        "vol2atlas alignFull state_em.json --tps -o out/em_in_ccf.zarr",
    ])

    para(doc, [("What would have to change vs Stage A.", {"bold": True})])
    bullet(doc, [
        ("Reference loader.  ", {"bold": True}),
        ("add load_zarr_reference(path) alongside load_ccf(atlas_name) and pivot the steps onto a Reference{array, voxel_um, name} dataclass. ~6 one-line edits across the step files.", {}),
    ])
    bullet(doc, [
        ("Voxel-size-aware UI.  ", {"bold": True}),
        ("Slider step (hardcoded 50.0 µm in step1_prealign.py:215) and preview stride (PREVIEW_MAX = 192 in step1_prealign.py:45) need to become functions of sample_voxel_um.", {}),
    ])
    bullet(doc, [
        ("Warp composition.  ", {"bold": True}),
        ("extend align_full.py's _warp_one_chunk to compose a TPS displacement per chunk on top of the existing rigid step. A B-spline component would be a later, separate change.", {}),
    ])

    para(doc, [("What is already in place (Stage A engine, reusable as-is).", {"bold": True})])
    bullet(doc, [
        ("Chunkwise rigid warp.  ", {"bold": True}),
        ("align_full.py already reads one output chunk's source-bbox at a time via dask_image.ndinterp and dispatches chunks with ThreadPoolExecutor. RAM is bounded by chunk size on µCT runs.", {}),
    ])
    bullet(doc, [
        ("State schema.  ", {"bold": True}),
        ("state.py already supports the Stage A fields; Stage B would add a single optional reference_zarr field. Two independent runs, chained only through the file path.", {}),
    ])

    para(doc, [("Open risks.", {"bold": True, "color": (0xB0, 0x30, 0x30)})])
    bullet(doc, [
        ("nm-scale UI ergonomics.  ", {"bold": True}),
        ("untested whether sliders parameterized in 0.004 µm steps are usable.", {}),
    ])
    bullet(doc, [
        ("OME-Zarr level selection on TB inputs.  ", {"bold": True}),
        ("Stage A typically starts from level 2; the right starting level for TB EM is unknown.", {}),
    ])
    bullet(doc, [
        ("TPS conditioning at nm.  ", {"bold": True}),
        ("how spread / coverage of nm landmark pairs affects TPS stability is not characterized.", {}),
    ])
    bullet(doc, [
        ("RAM extrapolation.  ", {"bold": True}),
        ("Stage A's per-chunk RAM bound has been measured at µCT chunk sizes only; TB EM chunks may behave differently.", {}),
    ])

    para(doc, [("Out of scope.", {"bold": True})])
    para(doc, [
        ("EM section assembly (TrakEM2 / render-python / ASAP) — input "
         "must already be a 3D OME-Zarr. Deformable registration on TB "
         "volumes directly is also out of scope; the design intent is that "
         "elastix / ANTs run on a downsampled level and the resulting "
         "B-spline composes into the chunkwise warp.", {}),
    ])

    out = Path(__file__).resolve().parent.parent / "vol2atlas_method.docx"
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
